# -*- coding: utf-8 -*-
"""
Created on Wed Jul  1 11:42:40 2020

@author: erich
"""
# Data structure using docx
# Document > Paragraph > Run Objects

import docx

# Create Document Obect file
d = docx.Document(r'C:\Users\erich\Desktop\Python\Sample Resume.docx')

# Call and print paragraph within the Document
p1 = d.paragraphs[0]
p2 = d.paragraphs[1]
print(p1.text+'\n'+p2.text)

# Call and print run objects within the Document
r1 = p1.runs[0]
r2 = p2.runs[0]
print(r1.text+'\n'+r2.text)

# Check for run  formatting
print(r1.bold==None) # Confirms the run is not bolded
print(r1.italic==None) # Confirms the run is not italic
print(r1.underline==None) # Confirms the run is not underline

# Assign formatting to runs
print('Q: is the run bold? '+str(r1.bold!=None))
r1.bold=True
print('Q: is the run bold? '+str(r1.bold!=None))
r1.bold=None
print('Q: is the run bold? '+str(r1.bold!=None))

# Save file
d.save(r'C:\Users\erich\Desktop\Python\Sample Resume 2.docx')

# Create new file
d3 = docx.Document() # Creates a doc
d3.add_paragraph(r'hello this is a text doc') # Add paragraph
p3 = d3.paragraphs[0] 
p3.add_run('. This is a new run within the paragraph') # Add run
r3 = p3.runs[1]
r3.bold=True # Bold the new run
d3.save(r'C:\Users\erich\Desktop\Python\Sample doc.docx') #Save


